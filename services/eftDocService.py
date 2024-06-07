#services/eftDocService.py

import docx
from docx.oxml.ns import nsmap

def get_element_value(element, tag, attribute='{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'):
    """Helper to fetch values with a given tag from an element."""
    found = element.find(f'.//w:{tag}', namespaces=nsmap)
    return found.get(attribute) if found is not None else None

def extract_table_data(table):
    """Extract data from table and return a list of rows with cell data."""
    table_data = []
    for row in table.findall('.//w:tr', namespaces=nsmap):
        row_data = []
        for cell in row.findall('.//w:tc', namespaces=nsmap):
            cell_text = ''
            for paragraph in cell.findall('.//w:p', namespaces=nsmap):
                for run in paragraph.findall('.//w:r', namespaces=nsmap):
                    for text in run.findall('.//w:t', namespaces=nsmap):
                        if text.text:
                            cell_text += text.text
            row_data.append(cell_text.strip())
        table_data.append(row_data)
    return table_data

def extract_content_control_data(element):
    """Extract data based on the type of content control."""
    title = get_element_value(element, 'alias')
    if title == "InterfacesGrid":
        table = element.find('.//w:tbl', namespaces=nsmap)
        return (title, extract_table_data(table)) if table is not None else (title, [])
    dropdown = element.find('.//w:dropDownList', namespaces=nsmap)
    if dropdown:
        last_value = get_element_value(dropdown, 'lastValue')
        return (title, last_value if last_value is not None else "No Selection")
    text_content = get_element_value(element, 't')
    return (title, text_content if text_content is not None else "No Text")

def extract_eft_data(doc_path):
    doc = docx.Document(doc_path)
    data = {}
    for element in doc.element.xpath('.//w:sdt'):
        title, content = extract_content_control_data(element)
        data[title] = content
    return data

def update_text(element, text):
    """Update or create text element within the specified element."""
    if isinstance(text, list):
        # Handle lists specially, assume it's for tables or similar structures
        update_table(element, text)
    else:
        text_element = element.find('.//w:t', namespaces=nsmap)
        if text_element is not None:
            text_element.text = text
        else:
            create_new_text(element, text)

def update_table(element, table_data):
    """Update table data within the specified element."""
    table = element.find('.//w:tbl', namespaces=nsmap)
    if table is not None:
        for row_index, row in enumerate(table.findall('.//w:tr', namespaces=nsmap)):
            if row_index < len(table_data):
                for cell_index, cell in enumerate(row.findall('.//w:tc', namespaces=nsmap)):
                    if cell_index < len(table_data[row_index]):
                        cell_text = table_data[row_index][cell_index]
                        paragraph = cell.find('.//w:p', namespaces=nsmap)
                        if paragraph is None:
                            paragraph = docx.oxml.shared.OxmlElement('w:p')
                            cell.append(paragraph)
                        run = paragraph.find('.//w:r', namespaces=nsmap)
                        if run is None:
                            run = docx.oxml.shared.OxmlElement('w:r')
                            paragraph.append(run)
                        text_element = run.find('.//w:t', namespaces=nsmap)
                        if text_element is None:
                            text_element = docx.oxml.shared.OxmlElement('w:t')
                            run.append(text_element)
                        text_element.text = cell_text

def create_new_text(element, text):
    """Create new text element if no existing text element is found."""
    new_text_element = docx.oxml.shared.OxmlElement('w:t')
    new_text_element.text = text
    new_run_element = docx.oxml.shared.OxmlElement('w:r')
    new_run_element.append(new_text_element)
    new_paragraph_element = docx.oxml.shared.OxmlElement('w:p')
    new_paragraph_element.append(new_run_element)
    element.append(new_paragraph_element)

def write_eft_data(doc_path, data):
    try:
        doc = docx.Document(doc_path)
        for element in doc.element.xpath('.//w:sdt'):
            title = get_element_value(element, 'alias')
            if title in data:
                update_text(element, data[title])
        doc.save(doc_path)
        print("Data written successfully to the document.")
    except PermissionError:
        print("Failed to write data. The document is currently open. Please close it and try again.")

def save_updated_eft_doc(doc_path, updated_doc_path):
    try:
        doc = docx.Document(doc_path)
        doc.save(updated_doc_path)
        print("Updated EFT document saved successfully.")
    except PermissionError:
        print("Failed to save the updated document. Please close it and try again.")

# Example usage
doc_path = '/Users/tomenns/Downloads/EFT Design Document Template Final.docx'

# Extract data from the document

eft_data = extract_eft_data(doc_path)
print("Extracted data:")
print(eft_data)

# Modify the extracted data this is for testing only and will be removed in final version.
eft_data['CaseNum'] = '126657745745745'
eft_data['QuoteNum'] = 'ABC123'
eft_data['Company'] = 'asdfasdfasdf'
eft_data['Status'] = '99.99'
eft_data['Abstract'] = 'TitleGOesHERERERERE'
eft_data['Database'] = 'TestDB'
eft_data['TransactionType'] = 'Lockbox'
eft_data['BankName'] = 'Bank of America'
eft_data['InterfacesGrid'][1][0] = 'Bank osdfrica'
eft_data['InterfacesGrid'][1][1] = 'Asdf'
eft_data['InterfacesGrid'][1][2] = 'Losdfox'
eft_data['InterfacesGrid'][1][3] = 'Spesdfdf'

# Write the modified data back to the document
write_eft_data(doc_path, eft_data)