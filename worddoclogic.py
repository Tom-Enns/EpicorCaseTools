import re
from docx import Document
import zipfile
from xml.etree import ElementTree as ET

def extract_fields_revised(doc_path):
    # Load the document
    doc = Document(doc_path)

    # Define the namespaces
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
        'ext': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'
    }

    # Extract content controls
    content_controls = {}
    for sdt in doc._element.findall('.//w:sdt', namespaces):
        tag = sdt.find('./w:sdtPr/w:tag', namespaces).get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
        content = "".join([t.text for t in sdt.findall('.//w:t', namespaces)])
        if tag and content:
            content_controls[tag] = content.strip()

    # Extract core property: Author
    author = doc.core_properties.author

    with zipfile.ZipFile(doc_path, 'r') as docx_zip:
        # Extract custom property: Company (Customer)
        app_xml_content = docx_zip.read('docProps/app.xml').decode()
        root = ET.fromstring(app_xml_content)
        company_element = root.find('ext:Company', namespaces)
        company = company_element.text if company_element is not None else ""

        # Extract title from customXml/item1.xml
        custom_xml_content = docx_zip.read('customXml/item1.xml').decode()
        title_start_index = custom_xml_content.find("<Abstract>")
        title_end_index = custom_xml_content.find("</Abstract>")
        title = custom_xml_content[title_start_index + 10:title_end_index]

        # Extract revision from docProps/core.xml
        core_xml_content = docx_zip.read('docProps/core.xml').decode()
        revision_start_index = core_xml_content.find("<cp:contentStatus>")
        revision_end_index = core_xml_content.find("</cp:contentStatus>")
        revision = core_xml_content[revision_start_index + 18:revision_end_index]

    # Extracted data
    extracted_data = {
        "Title": title,
        "Case": content_controls.get("CaseNum", ""),
        "Quote": content_controls.get("QuoteNum", ""),
        "Database": content_controls.get("Database", ""),
        "Author": author,
        "Customer": company,
        "Revision": revision
    }

    return extracted_data

def extract_need_using_docx(doc_path):
    doc = Document(doc_path)
    paragraphs = [p.text for p in doc.paragraphs]

    # Find the paragraph containing "What is your need?"
    need_start_index = next((i for i, p in enumerate(paragraphs) if "What is your need?" in p), None)
    
    if need_start_index is None:
        return None

    # Extract the subsequent paragraphs until reaching the next section
    need_content = []
    for para in paragraphs[need_start_index + 1:]:
        if "Response" in para:
            break
        need_content.append(para)

    return " ".join(need_content).strip()

def extract_all_fields(doc_path):
    # Extract main fields
    main_fields = extract_fields_revised(doc_path)
    
    # Extract "What is your need?" answer
    need_answer = extract_need_using_docx(doc_path)
    
    # Refine the extracted content to remove the placeholder example
    refined_need_content = need_answer.split("2 Changes needed:")[1].split("Do you have technical requirements?")[0].strip()
    
    # Add the refined need content to the main fields dictionary
    main_fields["Need"] = refined_need_content
    
    return main_fields

# Example use:
# data = extract_all_fields("/path/to/your/Design - Case 3059 V1 OLD2.docx")
# print(data)

